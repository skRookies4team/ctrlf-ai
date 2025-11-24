import pdfplumber
import docx
import easyocr
from bs4 import BeautifulSoup
from lxml import etree

def detect_type(file):
    name = file.name.lower()

    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith(".txt"):
        return "txt"
    if name.endswith(".md"):
        return "md"
    if name.endswith(".html") or name.endswith(".htm"):
        return "html"
    if name.endswith(".hwp"):
        return "hwp"
    if name.endswith(".hwpx"):
        return "hwpx"
    if name.endswith((".png", ".jpg", ".jpeg")):
        return "image"

    return "unknown"
def load_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            text += (page_text or "") + "\n"
    return text


def load_docx(file):
    doc = docx.Document(file)
    return "\n".join([p.text for p in doc.paragraphs])


def load_txt(file):
    return file.read().decode("utf-8", errors="ignore")


def load_html(file):
    html = file.read().decode("utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(" ", strip=True)


def load_image(file):
    reader = easyocr.Reader(["ko", "en"])
    result = reader.readtext(file.read(), detail=0)
    return "\n".join(result)


def load_hwp(file):
    import pyhwp
    doc = pyhwp.HWPDocument(file)
    doc.convert_text()
    return doc.body_text


def load_hwpx(file):
    xml = etree.parse(file)
    return "\n".join(
        xml.xpath("//w:t/text()", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    )
def load_document(file):
    doc_type = detect_type(file)

    if doc_type == "pdf":
        return load_pdf(file)
    if doc_type == "docx":
        return load_docx(file)
    if doc_type == "txt":
        return load_txt(file)
    if doc_type == "md":
        return load_txt(file)
    if doc_type == "html":
        return load_html(file)
    if doc_type == "image":
        return load_image(file)
    if doc_type == "hwp":
        return load_hwp(file)
    if doc_type == "hwpx":
        return load_hwpx(file)

    return ""
